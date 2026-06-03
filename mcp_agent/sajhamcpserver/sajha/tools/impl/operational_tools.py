"""
Operational Tools Suite — pdf_read, md_save, md_to_docx, search_files, fill_template, list_versions
"""
import io
import os, re, json, shutil, mimetypes
from datetime import datetime, timezone
from pathlib import Path
from sajha.tools.base_mcp_tool import BaseMCPTool
from sajha.core.properties_configurator import PropertiesConfigurator
from sajha.storage import storage
from sajha.path_resolver import resolve as path_resolve
from sajha.data_context import get_data_layers


def _get_worker_ctx():
    try:
        from flask import g as _g
        return getattr(_g, 'worker_ctx', {}) or {}
    except RuntimeError:
        return {}


def _props():
    return PropertiesConfigurator()

def _resolve(path_str):
    """Resolve a path relative to the sajhamcpserver working directory (CWD)."""
    p = Path(path_str)
    if p.is_absolute():
        return p.resolve()
    return (Path.cwd() / p).resolve()

def _domain_root():
    """Return domain data root via data_context (consistent with all other tools)."""
    layers = get_data_layers('domain_data')
    if layers:
        return Path(layers[0][1].rstrip('/')).resolve()
    return _resolve(_props().get('data.domain_data.dir', './data/domain_data'))

def _my_data_root():
    """Return my-data root via data_context (REQ-DD-02 / REQ-MD-01)."""
    layers = get_data_layers('my_data')
    if layers:
        return Path(layers[0][1].rstrip('/')).resolve()
    return _resolve(_props().get('data.my_data.dir', './data/uploads'))

def _common_root():
    """Return common/shared-library root via data_context."""
    layers = get_data_layers('common')
    if layers:
        return Path(layers[0][1].rstrip('/')).resolve()
    return _resolve(_props().get('data.common_data.dir', './data/common'))

def _templates_dir():
    """Return templates dir derived from domain_data root."""
    layers = get_data_layers('domain_data')
    if layers:
        return Path(layers[0][1].rstrip('/') + '/templates').resolve()
    return _resolve(_props().get('data.templates_dir', './data/domain_data/templates'))

class _S3PseudoPath:
    """Minimal pathlib.Path-compatible wrapper for s3:// URIs.
    Exposes .name, .stem, .suffix, .parent, .exists(), and str() so that
    callers written for local paths work transparently in S3 mode.
    """
    def __init__(self, uri: str):
        self._uri = uri
        self._basename = os.path.basename(uri.rstrip('/'))
        base, ext = os.path.splitext(self._basename)
        self._stem = base
        self._suffix = ext

    @property
    def name(self) -> str:
        return self._basename

    @property
    def stem(self) -> str:
        return self._stem

    @property
    def suffix(self) -> str:
        return self._suffix

    @property
    def parent(self) -> '_S3PseudoPath':
        parent_uri = self._uri[:self._uri.rfind('/')]
        return _S3PseudoPath(parent_uri)

    def with_suffix(self, suffix: str) -> '_S3PseudoPath':
        new_uri = self._uri[:self._uri.rfind('.')] + suffix if '.' in os.path.basename(self._uri) else self._uri + suffix
        return _S3PseudoPath(new_uri)

    def exists(self) -> bool:
        return storage.exists(self._uri)

    def __str__(self) -> str:
        return self._uri

    def __fspath__(self) -> str:
        return self._uri


def _find_file(name, *search_roots):
    """Search for a bare filename across the given roots (and one level of subdirs).
    Returns a Path if found, else None. Used so tools accept filenames without full paths.
    """
    import os as _os
    bare = _os.path.basename(name)
    for root in search_roots:
        root = Path(root).resolve()
        # Direct hit
        candidate = root / bare
        if candidate.exists():
            return candidate
        # One-level subdirectory search
        try:
            for sub in root.iterdir():
                if sub.is_dir():
                    c = sub / bare
                    if c.exists():
                        return c
        except (PermissionError, OSError):
            pass
    return None


def _safe_path(path_str, *allowed_roots):
    """Return a Path (or _S3PseudoPath) if within one of the allowed roots, else None.
    S3 mode: s3:// URIs bypass local-path validation; bucket-level IAM is the boundary.
    """
    if str(path_str).startswith('s3://'):
        return _S3PseudoPath(str(path_str))
    try:
        p = Path(path_str).resolve()
    except Exception:
        return None
    for root in allowed_roots:
        root = Path(root).resolve()
        try:
            p.relative_to(root)
            return p
        except ValueError:
            pass
    return None


# ─────────────────────────── pdf_read ────────────────────────────────────────

class PdfReadTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "pdf_read",
            "description": (
                "Extract text and tables from a local PDF file. "
                "Pass the absolute file_path from list_data_files or list_uploaded_files. "
                "Use pages='1-5' for large PDFs to limit output size."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path or bare filename of a .pdf file (e.g. 'report.pdf'). Use the path from list_uploaded_files, or just the filename — the tool will search all data layers automatically."},
                "pages": {"type": "string", "description": "Page range: '1', '1-5', or 'all' (default)."},
                "extract_tables": {"type": "boolean", "description": "Extract tables. Default: true."},
                "max_chars": {"type": "integer", "description": "Truncate text at this char count. Default: 50000."},
                "heading": {
                    "type": "string",
                    "description": (
                        "For large PDFs: extract only the section under this heading. "
                        "Detected using font-size analysis — larger/bolder text is treated as a heading. "
                        "Case-insensitive partial match. E.g. 'Market Risk', 'Item 7', 'Executive Summary'. "
                        "If not found, returns available_headings[] detected in the document."
                    ),
                },
            },
            "required": ["file_path"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        file_path = arguments.get("file_path", "")
        pages_arg = arguments.get("pages", "all")
        extract_tables = arguments.get("extract_tables", True)
        max_chars = arguments.get("max_chars", 50000)
        heading = (arguments.get("heading") or "").strip()

        safe = _safe_path(file_path, _domain_root(), _my_data_root(), _common_root())
        if not safe or not storage.exists(str(safe)):
            # Fallback: treat as bare filename and search all roots
            safe = _find_file(file_path, _domain_root(), _my_data_root(), _common_root())
        if not safe or not storage.exists(str(safe)):
            return {"error": f"File not found or access denied: {file_path}"}
        if safe.suffix.lower() != ".pdf":
            return {"error": "pdf_read only accepts .pdf files."}

        try:
            import fitz  # PyMuPDF
        except ImportError:
            try:
                import pdfplumber as _pl
                return self._read_pdfplumber(safe, pages_arg, max_chars)
            except ImportError:
                return {"error": "Neither PyMuPDF nor pdfplumber is installed."}

        try:
            raw = storage.read_bytes(str(safe))
            doc = fitz.open(stream=io.BytesIO(raw), filetype='pdf')
        except Exception as e:
            return {"error": f"PDF could not be parsed: {e}"}

        if doc.is_encrypted:
            return {"error": "PDF is encrypted. Cannot extract without password."}

        total_pages = len(doc)
        page_indices = self._parse_pages(pages_arg, total_pages)

        # ── Heading extraction mode ──────────────────────────────────────────
        if heading:
            return self._extract_pdf_section(doc, page_indices, heading, safe, total_pages)

        # ── Standard full read ───────────────────────────────────────────────
        text_parts = []
        tables = []
        for i in page_indices:
            page = doc[i]
            text_parts.append(page.get_text())
            if extract_tables:
                for tab in page.find_tables():
                    headers = tab.header.names if hasattr(tab, 'header') else []
                    rows = tab.extract()
                    if rows:
                        if not headers and rows:
                            headers = rows[0]
                            rows = rows[1:]
                        tables.append({"page": i + 1, "headers": headers, "rows": rows})

        full_text = "\n".join(text_parts)
        truncated = len(full_text) > max_chars
        if truncated:
            full_text = full_text[:max_chars]

        warning = None
        if not full_text.strip():
            warning = "No text layer detected. PDF may be image-only."

        result = {
            "filename": safe.name,
            "pages_extracted": len(page_indices),
            "total_pages": total_pages,
            "char_count": len(full_text),
            "truncated": truncated,
            "text": full_text,
            "_source": str(safe),
        }
        if extract_tables:
            result["tables"] = tables
        if warning:
            result["warning"] = warning
        return result

    def _detect_pdf_headings(self, doc, page_indices):
        """
        Detect heading-like lines using font-size analysis via PyMuPDF.

        Strategy:
          1. Collect every text line with its max font size and bold flag, per page.
          2. Compute the most-common (body) font size via Counter.
          3. A line is a heading if font_size > body_size * 1.15 OR
             (bold AND short AND >= body_size).
          4. Exclude: TOC leader lines (`...`), pure numbers, very long lines.
          5. Exclude running headers/footers — text that appears on ≥30% of pages.
          6. Allow multiple occurrences of the same heading (e.g. TOC + body);
             the caller will pick the occurrence with the longest following content.

        Returns (headings, full_text):
          headings — list of {'text', 'page', 'char_offset'} dicts
          full_text — concatenated page texts (joined by newline) for slicing
        """
        from collections import Counter

        # Per-page line collection
        page_line_items = {}  # page_idx -> [(font_size, is_bold, text), ...]
        for page_idx in page_indices:
            items = []
            page = doc[page_idx]
            d = page.get_text("dict")
            for block in d.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    parts, max_size, bold = [], 0.0, False
                    for span in line.get("spans", []):
                        t = span.get("text", "")
                        if t.strip():
                            parts.append(t)
                        sz = span.get("size", 0.0)
                        if sz > max_size:
                            max_size = sz
                        if span.get("flags", 0) & 16:
                            bold = True
                    line_text = "".join(parts).strip()
                    if line_text:
                        items.append((max_size, bold, line_text))
            page_line_items[page_idx] = items

        all_line_items = [(p + 1, sz, bold, txt)
                         for p, items in page_line_items.items()
                         for sz, bold, txt in items]

        if not all_line_items:
            return [], ""

        # Body font size = most common rounded size
        body_size = Counter(round(it[1]) for it in all_line_items).most_common(1)[0][0]

        # Detect running headers/footers: normalized text appearing on ≥30% of pages
        n_pages = len(page_indices)
        text_page_count: Counter = Counter()
        for p, sz, bold, txt in all_line_items:
            text_page_count[re.sub(r'\s+', ' ', txt.lower()[:40])] += 1
        running_threshold = max(2, int(n_pages * 0.30))
        running_texts = {k for k, v in text_page_count.items() if v >= running_threshold}

        # Build full text
        full_text = "\n".join(doc[p].get_text() for p in page_indices)

        headings = []
        search_from = 0  # advance through full_text to find offsets in order
        for page, size, bold, text in all_line_items:
            norm_key = re.sub(r'\s+', ' ', text.lower()[:40])
            is_large   = size > body_size * 1.15
            is_bold_hdr = bold and len(text) < 120 and size >= body_size * 0.95
            is_toc_line = bool(re.search(r'\.{3,}', text))
            is_numeric  = bool(re.match(r'^[\d\s\.\,\-\$\%]+$', text))
            is_running  = norm_key in running_texts
            too_long    = len(text) > 200
            too_short   = len(text) < 3
            if (is_large or is_bold_hdr) and not any(
                    [is_toc_line, is_numeric, is_running, too_long, too_short]):
                # Locate this text in full_text (search forward from last offset)
                pos = full_text.lower().find(text.lower()[:40], search_from)
                if pos == -1:
                    pos = full_text.lower().find(text.lower()[:40])
                if pos != -1:
                    search_from = pos + 1
                    headings.append({"text": text, "page": page, "char_offset": pos})

        return headings, full_text

    def _extract_pdf_section(self, doc, page_indices, heading_query, safe, total_pages):
        """
        Extract body text between the matched heading and the next detected heading
        using character offsets into the full concatenated document text.
        """
        result = self._detect_pdf_headings(doc, page_indices)
        if isinstance(result, tuple):
            headings, full_text = result
        else:
            headings, full_text = result, ""

        if not headings:
            return {
                "error": (
                    "No headings detected in this PDF. "
                    "The document may use image-only text or non-standard formatting. "
                    "Try pdf_read without heading= and use pages= to narrow the range."
                ),
                "filename": safe.name,
                "_source": str(safe),
            }

        query = heading_query.strip().lower()
        match_idx = None
        for i, h in enumerate(headings):
            if query in h["text"].lower():
                match_idx = i
                break

        available = [h["text"] for h in headings]

        if match_idx is None:
            return {
                "error": f'Heading "{heading_query}" not found in {safe.name}',
                "available_headings": available[:50],
                "filename": safe.name,
                "total_pages": total_pages,
                "_source": str(safe),
            }

        # Find the best match — PDFs often repeat headings in the TOC (short content)
        # before the real body section (long content). Also, the TOC and body versions
        # of the same heading may differ slightly (e.g. em-dash vs hyphen).
        # Strategy: collect ALL headings from the list that match the query, compute
        # the content length from each to the next heading, pick the longest.
        all_heading_offsets = sorted(h["char_offset"] for h in headings)

        matching_headings = [(i, h) for i, h in enumerate(headings)
                             if query in h["text"].lower()]

        best_text = ""
        best_matched_text = headings[match_idx]["text"]
        best_start_page = headings[match_idx]["page"]

        for _, candidate_h in matching_headings:
            start_off = candidate_h["char_offset"]
            # Next heading boundary after this one
            next_boundary = len(full_text)
            for hoff in all_heading_offsets:
                if hoff > start_off + len(candidate_h["text"]):
                    next_boundary = hoff
                    break
            candidate_text = full_text[start_off:next_boundary].strip()
            if len(candidate_text) > len(best_text):
                best_text = candidate_text
                best_matched_text = candidate_h["text"]
                best_start_page = candidate_h["page"]

        section_text = best_text
        max_chars = 60_000
        truncated = len(section_text) > max_chars
        if truncated:
            section_text = section_text[:max_chars]

        return {
            "filename": safe.name,
            "matched_heading": best_matched_text,
            "start_page": best_start_page,
            "total_pages": total_pages,
            "char_count": len(section_text),
            "truncated": truncated,
            "content": section_text,
            "available_headings": available[:50],
            "_source": str(safe),
        }

    def _read_pdfplumber(self, safe, pages_arg, max_chars):
        """Fallback PDF reader using pdfplumber when PyMuPDF is unavailable."""
        import pdfplumber
        try:
            raw = storage.read_bytes(str(safe))
            pdf = pdfplumber.open(io.BytesIO(raw))
        except Exception as e:
            return {"error": f"PDF could not be parsed: {e}"}

        total_pages = len(pdf.pages)
        page_indices = self._parse_pages(pages_arg, total_pages)

        text_parts = []
        tables = []
        for i in page_indices:
            page = pdf.pages[i]
            text = page.extract_text() or ""
            text_parts.append(text)
            for tab in (page.extract_tables() or []):
                if tab:
                    headers = [str(c) if c else "" for c in tab[0]]
                    rows = [[str(c) if c else "" for c in row] for row in tab[1:]]
                    tables.append({"page": i + 1, "headers": headers, "rows": rows})

        pdf.close()
        full_text = "\n".join(text_parts)
        truncated = len(full_text) > max_chars
        if truncated:
            full_text = full_text[:max_chars]

        result = {
            "filename": safe.name,
            "pages_extracted": len(page_indices),
            "total_pages": total_pages,
            "char_count": len(full_text),
            "truncated": truncated,
            "text": full_text,
            "_source": str(safe),
        }
        result["tables"] = tables
        if not full_text.strip():
            result["warning"] = "No text layer detected. PDF may be image-only."
        return result

    def _parse_pages(self, pages_arg, total):
        if pages_arg == "all" or not pages_arg:
            return list(range(total))
        m = re.match(r"^(\d+)-(\d+)$", pages_arg.strip())
        if m:
            a, b = int(m.group(1)) - 1, int(m.group(2)) - 1
            return list(range(max(0, a), min(total, b + 1)))
        m2 = re.match(r"^(\d+)$", pages_arg.strip())
        if m2:
            i = int(m2.group(1)) - 1
            return [i] if 0 <= i < total else []
        return list(range(total))


# ─────────────────────────── md_save ─────────────────────────────────────────

class MdSaveTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "md_save",
            "description": (
                "Save a Markdown string to my_data/ as a .md file with automatic versioning. "
                "If the file already exists, the old version is archived with a timestamp suffix. "
                "Use this to persist analysis outputs, filled templates, and canvas content."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "content": {"type": "string", "description": "Markdown content to write."},
                "filename": {"type": "string", "description": "Target filename including .md extension."},
                "subfolder": {"type": "string", "description": "Sub-folder within my_data/. Created if needed. Default: ''."},
                "versioning": {"type": "boolean", "description": "Archive existing file before writing. Default: true."},
                "overwrite": {"type": "boolean", "description": "Overwrite without archiving. Overrides versioning. Default: false."},
            },
            "required": ["content", "filename"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        content = arguments.get("content", "")
        filename = Path(arguments.get("filename", "output.md")).name
        if not filename.endswith(".md"):
            filename += ".md"
        subfolder = arguments.get("subfolder", "")
        versioning = arguments.get("versioning", True)
        overwrite = arguments.get("overwrite", False)

        root = _my_data_root()
        if subfolder:
            folder = root / subfolder
        else:
            folder = root
        folder.mkdir(parents=True, exist_ok=True)

        dest = folder / filename
        archived_as = None

        if storage.exists(str(dest)) and not overwrite:
            if versioning:
                ts = datetime.now(tz=timezone.utc).strftime("%Y-%m-%d_%H%M%S")
                stem = dest.stem
                archive_name = f"{stem}_{ts}.md"
                archive_path = folder / archive_name
                storage.copy(str(dest), str(archive_path))
                storage.delete(str(dest))
                archived_as = archive_name

        storage.write_text(str(dest), content, encoding="utf-8")
        size_bytes = storage.get_size(str(dest)) or len(content.encode("utf-8"))

        return {
            "path": str(dest),
            "filename": filename,
            "subfolder": subfolder,
            "size_bytes": size_bytes,
            "versioned": archived_as is not None,
            "archived_as": archived_as,
            "written_at": datetime.now(tz=timezone.utc).isoformat(),
            "_source": str(dest),
        }


# ─────────────────────────── md_to_docx ──────────────────────────────────────

class MdToDocxTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "md_to_docx",
            "description": (
                "Convert a .md file to a formatted .docx Word document. "
                "Handles headings, tables, bullet lists, bold/italic, code blocks, and YAML frontmatter (omitted). "
                "Output is saved alongside the source file or to output_path."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Absolute path to a .md file."},
                "output_path": {"type": "string", "description": "Where to save .docx. Defaults to alongside source."},
                "style": {"type": "string", "enum": ["standard", "minimal", "report"], "description": "Style preset. Default: standard."},
                "include_toc": {"type": "boolean", "description": "Insert Table of Contents. Default: false."},
            },
            "required": ["file_path"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        file_path = arguments.get("file_path", "")
        output_path = arguments.get("output_path")
        include_toc = arguments.get("include_toc", False)

        safe = _safe_path(file_path, _domain_root(), _my_data_root(), _common_root())
        if not safe or not storage.exists(str(safe)):
            return {"error": f"File not found or access denied: {file_path}"}
        if safe.suffix.lower() != ".md":
            return {"error": "md_to_docx only accepts .md files."}

        text = storage.read_text(str(safe), encoding="utf-8")
        # Strip YAML frontmatter
        text = re.sub(r"^---\n.*?\n---\n", "", text, flags=re.DOTALL)

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            return {"error": "python-docx is not installed."}

        doc = Document()
        # Set narrow margins
        for section in doc.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        def set_heading_style(para, level):
            colors = {1: RGBColor(0x1e, 0x3a, 0x5f), 2: RGBColor(0x1e, 0x3a, 0x5f), 3: RGBColor(0x1e, 0x6e, 0xb0)}
            para.style = f'Heading {level}'
            for run in para.runs:
                run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))

        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Headings
            m = re.match(r"^(#{1,6})\s+(.*)", line)
            if m:
                level = min(len(m.group(1)), 3)
                para = doc.add_heading(m.group(2).strip(), level=level)
                i += 1
                continue

            # HR
            if re.match(r"^---+$", line.strip()):
                para = doc.add_paragraph()
                i += 1
                continue

            # Code block
            if line.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing ```
                para = doc.add_paragraph("\n".join(code_lines))
                para.style = "Normal"
                for run in para.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                continue

            # Table (pipe syntax)
            if "|" in line and line.strip().startswith("|"):
                table_rows = []
                while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                    row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    # Skip separator row
                    if not all(re.match(r"^[-:]+$", c) for c in row if c):
                        table_rows.append(row)
                    i += 1
                if table_rows:
                    max_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=max_cols)
                    table.style = "Table Grid"
                    for ri, row in enumerate(table_rows):
                        for ci, cell_text in enumerate(row):
                            if ci < max_cols:
                                cell = table.rows[ri].cells[ci]
                                cell.text = cell_text
                                if ri == 0:
                                    for run in cell.paragraphs[0].runs:
                                        run.font.bold = True
                continue

            # Bullet list
            if re.match(r"^[\-\*\+]\s+", line):
                text_part = re.sub(r"^[\-\*\+]\s+", "", line)
                para = doc.add_paragraph(style="List Bullet")
                self._add_inline(para, text_part)
                i += 1
                continue

            # Numbered list
            if re.match(r"^\d+\.\s+", line):
                text_part = re.sub(r"^\d+\.\s+", "", line)
                para = doc.add_paragraph(style="List Number")
                self._add_inline(para, text_part)
                i += 1
                continue

            # Blockquote
            if line.startswith("> "):
                para = doc.add_paragraph(line[2:])
                para.paragraph_format.left_indent = Inches(0.5)
                i += 1
                continue

            # Empty line
            if not line.strip():
                i += 1
                continue

            # Normal paragraph
            para = doc.add_paragraph()
            self._add_inline(para, line)
            i += 1

        if output_path:
            out = Path(output_path)
        else:
            out = safe.with_suffix(".docx")

        import io as _io
        _buf = _io.BytesIO()
        doc.save(_buf)
        _docx_bytes = _buf.getvalue()
        storage.write_bytes(str(out), _docx_bytes)
        return {
            "source_md": str(safe),
            "output_docx": str(out),
            "size_bytes": len(_docx_bytes),
            "style": arguments.get("style", "standard"),
            "toc_included": include_toc,
            "_source": str(out),
        }

    def _add_inline(self, para, text):
        """Add text with **bold** and *italic* inline formatting."""
        # Split on bold/italic markers
        pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
        parts = re.split(pattern, text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = para.add_run(part[1:-1])
                run.font.name = "Courier New"
                from docx.shared import Pt
                run.font.size = Pt(10)
            else:
                para.add_run(part)


# ─────────────────────────── search_files ────────────────────────────────────

class SearchFilesTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "search_files",
            "description": (
                "Full-text keyword search across domain_data/, my_data/, and common/ (Shared Library) files. "
                "Returns file path and excerpts with highlighted matches for each hit. "
                "Supports pdf, docx, xlsx, csv, md, txt, json. "
                "Use to find documents containing a keyword or exact phrase before processing them."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Search term. Case-insensitive. Use quotes for exact phrase."},
                "section": {"type": "string", "enum": ["domain_data", "my_data", "common", "all"], "description": "Limit search scope. Default: all (includes domain_data, my_data, and common/Shared Library)."},
                "file_type": {"type": "string", "description": "Filter by extension: pdf|docx|xlsx|csv|md|txt|json|all. Default: all."},
                "folder": {"type": "string", "description": "Limit to a sub-folder name (partial match)."},
                "max_results": {"type": "integer", "description": "Max file matches. Default: 20."},
                "excerpt_chars": {"type": "integer", "description": "Context chars around each match. Default: 200."},
            },
            "required": ["query"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        query = arguments.get("query", "")
        section = arguments.get("section", "all")
        file_type = arguments.get("file_type", "all")
        folder_filter = arguments.get("folder", "")
        max_results = arguments.get("max_results", 20)
        excerpt_chars = arguments.get("excerpt_chars", 200)

        # Build search pattern
        exact = re.match(r'^"(.+)"$', query)
        if exact:
            pattern = re.compile(re.escape(exact.group(1)), re.IGNORECASE)
        else:
            pattern = re.compile(re.escape(query), re.IGNORECASE)

        # Collect candidate root paths.
        # Use raw worker context strings (not resolved absolute paths) so that
        # storage._key() normalises them correctly in S3 mode.
        try:
            from flask import g as _g
            _dd = getattr(_g, 'worker_data_root', None) or ''
            _cd = getattr(_g, 'worker_common_root', None) or ''
            _md = getattr(_g, 'worker_my_data_root', None) or ''
        except RuntimeError:
            _dd = _cd = _md = ''
        _dd = _dd or _props().get('data.domain_data.dir', './data/domain_data')
        _cd = _cd or _props().get('data.common_data.dir', './data/common')
        _md = _md or _props().get('data.my_data.dir', './data/uploads')

        roots = []
        if section in ("domain_data", "all") and _dd:
            roots.append(("domain_data", _dd.rstrip('/')))
        if section in ("my_data", "all") and _md:
            roots.append(("my_data", _md.rstrip('/')))
        if section in ("common", "all") and _cd:
            roots.append(("common", _cd.rstrip('/')))

        candidates = []
        for sec_name, root_str in roots:
            # Use storage.list_prefix — works for both local filesystem and S3
            try:
                rel_paths = storage.list_prefix(root_str)
            except Exception:
                rel_paths = []
            for rel_path in rel_paths:
                fname = os.path.basename(rel_path)
                if fname.startswith("."):
                    continue
                if folder_filter and folder_filter.lower() not in rel_path.lower():
                    continue
                ext = os.path.splitext(fname)[1].lower().lstrip(".")
                if file_type != "all" and ext != file_type:
                    continue
                if ext in ("parquet", "pq", "db", "wal", "pyc"):
                    continue
                # Build the full path string (works for both local and s3://)
                f_path = root_str + '/' + rel_path
                candidates.append((sec_name, f_path, fname, ext))

        results = []
        for sec_name, f_path, fname, ext in candidates:
            if len(results) >= max_results:
                break
            try:
                text = self._extract_text(f_path, ext)
            except Exception:
                continue
            if not text:
                continue
            matches = list(pattern.finditer(text))
            if not matches:
                continue
            excerpts = []
            for m in matches[:5]:
                start = max(0, m.start() - excerpt_chars // 2)
                end = min(len(text), m.end() + excerpt_chars // 2)
                snippet = text[start:end].replace("\n", " ").strip()
                # Highlight match
                snippet = pattern.sub(lambda x: f"[{x.group()}]", snippet)
                excerpts.append(snippet)
            results.append({
                "filename": fname,
                "path": f_path,
                "file_type": ext,
                "section": sec_name,
                "match_count": len(matches),
                "excerpts": excerpts,
                "_source": f_path,
            })

        return {
            "query": query,
            "total_matches": len(results),
            "results": results,
        }

    def _extract_text(self, path, ext):
        if ext in ("md", "txt", "csv", "tsv"):
            return storage.read_text(str(path), encoding="utf-8")
        if ext == "json":
            return storage.read_text(str(path), encoding="utf-8")
        if ext == "docx":
            from docx import Document
            # Use BytesIO so this works with S3 paths and local paths alike
            raw = storage.read_bytes(str(path))
            doc = Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        if ext in ("xlsx", "xls"):
            import openpyxl
            # Use BytesIO so this works with S3 paths and local paths alike
            raw = storage.read_bytes(str(path))
            wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    parts.append(" ".join(str(c) for c in row if c is not None))
            return "\n".join(parts)
        if ext == "pdf":
            try:
                import fitz
                raw = storage.read_bytes(str(path))
                doc = fitz.open(stream=io.BytesIO(raw), filetype='pdf')
                return "\n".join(doc[i].get_text() for i in range(len(doc)))
            except Exception:
                return ""
        return ""


# ─────────────────────────── fill_template ───────────────────────────────────

class FillTemplateTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "fill_template",
            "description": (
                "Read a .md template from domain_data/templates/, substitute {{placeholder}} tokens "
                "with values from a data dict, and save the filled output to my_data/ via md_save. "
                "Templates declare required placeholders in YAML frontmatter. "
                "Use list_data_files with folder='templates' to discover available templates."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "template_path": {"type": "string", "description": "Absolute path to a .md template in domain_data/templates/."},
                "data": {"type": "object", "description": "Key-value dict mapping placeholder names to replacement values."},
                "output_filename": {"type": "string", "description": "Output filename. Defaults to template stem + '_filled_' + timestamp + '.md'."},
                "output_subfolder": {"type": "string", "description": "Sub-folder within my_data/. Default: 'reports'."},
                "versioning": {"type": "boolean", "description": "Archive existing output file. Default: true."},
                "convert_to_docx": {"type": "boolean", "description": "Also convert output to .docx. Default: false."},
            },
            "required": ["template_path", "data"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        template_path = arguments.get("template_path", "")
        data = arguments.get("data", {})
        output_subfolder = arguments.get("output_subfolder", "reports")
        versioning = arguments.get("versioning", True)
        convert_to_docx = arguments.get("convert_to_docx", False)

        # Security: template must be within worker templates dir or common templates dir
        tmpl_dir = _templates_dir()
        common_tmpl_dir = _common_root() / "templates"
        safe = _safe_path(template_path, tmpl_dir)
        if not safe or not storage.exists(str(safe)):
            safe = _safe_path(template_path, common_tmpl_dir)
        if not safe or not storage.exists(str(safe)):
            # Fallback: bare filename search across templates dirs
            safe = _find_file(template_path, tmpl_dir, common_tmpl_dir)
        if not safe or not storage.exists(str(safe)):
            return {"error": "Template not found or access denied"}
        if safe.suffix.lower() != ".md":
            return {"error": "fill_template only accepts .md template files."}

        content = storage.read_text(str(safe), encoding="utf-8")

        # Strip and parse frontmatter
        fm_match = re.match(r"^---\n(.*?)\n---\n", content, re.DOTALL)
        declared_placeholders = []
        if fm_match:
            fm_text = fm_match.group(1)
            content = content[fm_match.end():]
            # Extract placeholders list from frontmatter
            pm = re.search(r"placeholders:\s*\[([^\]]*)\]", fm_text)
            if pm:
                declared_placeholders = [p.strip().strip("'\"") for p in pm.group(1).split(",") if p.strip()]

        # Substitute placeholders
        filled = content
        for key, value in data.items():
            filled = filled.replace("{{" + key + "}}", str(value))

        # Find missing placeholders (tokens still present)
        remaining = re.findall(r"\{\{(\w+)\}\}", filled)
        missing = list(set(remaining))

        # Determine output filename
        output_filename = arguments.get("output_filename")
        if not output_filename:
            ts = datetime.now(tz=timezone.utc).strftime("%Y%m%d%H%M%S")
            output_filename = f"{safe.stem}_filled_{ts}.md"

        # Save via MdSaveTool logic directly
        saver = MdSaveTool()
        save_result = saver.execute({
            "content": filled,
            "filename": output_filename,
            "subfolder": output_subfolder,
            "versioning": versioning,
        })

        result = {
            "template_used": safe.name,
            "output_md_path": save_result.get("path"),
            "output_docx_path": None,
            "placeholders_filled": len(data) - len(missing),
            "placeholders_missing": missing,
            "versioned": save_result.get("versioned", False),
            "archived_as": save_result.get("archived_as"),
            "_source": save_result.get("path"),
        }

        if convert_to_docx and save_result.get("path"):
            converter = MdToDocxTool()
            docx_result = converter.execute({"file_path": save_result["path"]})
            if "output_docx" in docx_result:
                result["output_docx_path"] = docx_result["output_docx"]

        return result


# ─────────────────────────── list_versions ───────────────────────────────────

class ListVersionsTool(BaseMCPTool):
    def __init__(self, config=None):
        cfg = {
            "name": "list_versions",
            "description": (
                "List all versioned copies of a file stem in my_data/. "
                "Returns canonical (latest) file and all archived versions sorted newest first. "
                "Use after md_save to inspect version history or select a previous version to restore."
            ),
            "version": "1.0.0",
            "enabled": True,
        }
        if config:
            cfg.update(config)
        super().__init__(cfg)

    def get_input_schema(self):
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Canonical filename (with extension). E.g. 'rbc_ccr_brief.md'."},
                "subfolder": {"type": "string", "description": "Sub-folder within my_data/ to search. Omit to search all of my_data/."},
            },
            "required": ["filename"],
        }

    def get_output_schema(self):
        return {"type": "object"}

    def execute(self, arguments):
        filename = arguments.get("filename", "")
        subfolder = arguments.get("subfolder", "")

        root = _my_data_root()
        if subfolder:
            search_root = root / subfolder
        else:
            search_root = root

        stem = Path(filename).stem
        ext = Path(filename).suffix

        # Pattern for versioned files: stem_YYYY-MM-DD_HHmmss.ext
        version_pattern = re.compile(
            r"^" + re.escape(stem) + r"_(\d{4}-\d{2}-\d{2}_\d{6})" + re.escape(ext) + r"$",
            re.IGNORECASE
        )

        canonical = None
        versions = []

        # Use storage.list_prefix so this works with S3 and local backends
        try:
            rel_paths = storage.list_prefix(str(search_root))
        except Exception:
            rel_paths = []
        for rel_path in rel_paths:
            fname = os.path.basename(rel_path)
            f_path = str(search_root).rstrip('/') + '/' + rel_path
            size = storage.get_size(f_path)
            mtime = datetime.now(tz=timezone.utc)  # S3 has no local mtime; use now as fallback

            if fname.lower() == filename.lower():
                canonical = {
                    "filename": fname,
                    "path": f_path,
                    "modified_at": mtime.isoformat(),
                    "size_bytes": size,
                }
            else:
                m = version_pattern.match(fname)
                if m:
                    versions.append({
                        "filename": fname,
                        "path": f_path,
                        "archived_at": m.group(1),
                        "size_bytes": size,
                        "_ts": m.group(1),
                    })

        # Sort versions newest first
        versions.sort(key=lambda x: x["_ts"], reverse=True)
        for v in versions:
            del v["_ts"]

        if not canonical and versions:
            canonical = dict(versions[0])
            canonical["note"] = "No canonical file found; showing most recent version."

        return {
            "canonical": canonical,
            "versions": versions,
            "version_count": len(versions),
            "_source": str(search_root),
        }

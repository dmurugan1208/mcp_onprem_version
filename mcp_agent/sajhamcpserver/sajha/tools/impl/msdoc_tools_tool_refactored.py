"""
Copyright All rights Reserved 2025-2030, Ashutosh Sinha, Email: ajsinha@gmail.com
Microsoft Document Tools MCP Tool Implementation - Refactored with Individual Tools
"""

import os
import io
import json
from typing import Dict, Any, List, Optional
from sajha.tools.base_mcp_tool import BaseMCPTool
from sajha.core.properties_configurator import PropertiesConfigurator
from sajha.storage import storage


# Section parameter definition reused across all tool schemas
_SECTION_PROP = {
    "type": "string",
    "enum": ["domain_data", "my_data", "common"],
    "description": (
        "Data layer: 'domain_data' (worker knowledge base, default), "
        "'my_data' (user uploads), 'common' (shared library)."
    ),
}


class MsDocBaseTool(BaseMCPTool):
    """
    Base class for Microsoft Document tools with shared functionality
    """

    def __init__(self, config: Dict = None):
        """Initialize MS Doc base tool"""
        super().__init__(config)
        # Fallback docs directory (used when flask g context is unavailable)
        self.docs_directory = (
            (config.get('docs_directory') if config else None)
            or PropertiesConfigurator().get('tool.msdoc.docs_directory', 'data/msdocs')
        )

    def _resolve_docs_dir(self, section: str = 'domain_data') -> str:
        """Return the primary docs directory for the given section, using per-request
        worker context (flask g) when available, falling back to config."""
        try:
            from flask import g as _g
            if section == 'my_data':
                r = getattr(_g, 'worker_my_data_root', None)
                if r:
                    return r.rstrip('/')
            elif section == 'common':
                r = getattr(_g, 'worker_common_root', None)
                if r:
                    return r.rstrip('/')
            else:  # domain_data
                r = getattr(_g, 'worker_data_root', None)
                if r:
                    return r.rstrip('/')
        except RuntimeError:
            pass
        return self.docs_directory

    def _candidate_paths(self, filename: str) -> list:
        """Return all candidate absolute paths to search for a file, in priority order.

        Searches across all data layers so the tool finds the file regardless of
        which folder it was uploaded to:
          1. domain_data root
          2. domain_data/msdocs  (legacy subfolder)
          3. my_data/{user_id}   (user private files)
          4. common              (shared library)
        """
        candidates = []
        try:
            from flask import g as _g
            # domain_data root (highest priority — where files are normally uploaded)
            r = getattr(_g, 'worker_data_root', None)
            if r:
                candidates.append(os.path.join(r.rstrip('/'), filename))
                candidates.append(os.path.join(r.rstrip('/'), 'msdocs', filename))
            # my_data
            r = getattr(_g, 'worker_my_data_root', None)
            if r:
                candidates.append(os.path.join(r.rstrip('/'), filename))
            # common
            r = getattr(_g, 'worker_common_root', None)
            if r:
                candidates.append(os.path.join(r.rstrip('/'), filename))
        except RuntimeError:
            pass
        # Fallback to config-based directory
        candidates.append(os.path.join(self.docs_directory, filename))
        return candidates

    def _get_file_path(self, filename: str, section: str = 'domain_data') -> str:
        """Get full absolute path for a file, searching across all data folders.

        If section is specified it is tried first; then all other folders are
        checked in order so the tool finds the file regardless of where it lives.
        Also does a recursive walk so files inside subfolders (e.g. 'bmo financials/')
        are found even when the caller passes just the bare filename.
        """
        # Try the requested section's primary dir first (handles subfolder-qualified names too)
        primary = os.path.join(self._resolve_docs_dir(section), filename)
        if storage.exists(primary):
            return primary
        # Fall back to searching all candidate paths
        for path in self._candidate_paths(filename):
            if storage.exists(path):
                return path
        # Last resort: recursive walk across domain_data for bare filenames
        # (handles "Suppq425.xlsx" when it lives in "domain_data/bmo financials/Suppq425.xlsx")
        bare = os.path.basename(filename)
        try:
            from flask import g as _g
            roots_to_walk = []
            for attr in ('worker_data_root', 'worker_my_data_root', 'worker_common_root'):
                r = getattr(_g, attr, None)
                if r:
                    roots_to_walk.append(r.rstrip('/'))
            for root in roots_to_walk:
                for rel_key in storage.list_prefix(root):
                    if os.path.basename(rel_key) == bare:
                        return os.path.join(root, rel_key)
        except (RuntimeError, Exception):
            pass
        # Nothing found — return primary path so callers get a meaningful "not found" error
        return primary

    def _extract_word_heading(self, doc, heading_query: str) -> tuple:
        """
        Extract the content block under a matching Word heading style.

        Walks all paragraphs, identifies those with a 'Heading N' style, finds
        the first one whose text contains heading_query (case-insensitive), then
        collects paragraphs (and table text) until the next heading at the same
        or higher level.

        Returns (text, matched_title, available_headings).
        text is None when heading_query is not found.
        """
        query = heading_query.strip().lower()

        # Walk the document body in order, collecting paragraphs and table cells
        from docx.table import Table as _DocxTable
        from docx.text.paragraph import Paragraph as _DocxPara

        body_children = list(doc.element.body)

        # --- Pass 1: index all heading paragraphs ---
        heading_map = []  # list of (child_index, level, text)
        for ci, child in enumerate(body_children):
            # Paragraph elements
            if child.tag.endswith('}p'):
                para = _DocxPara(child, doc)
                style_name = para.style.name if para.style else ''
                if style_name.startswith('Heading'):
                    text = para.text.strip()
                    if not text:
                        continue
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    heading_map.append((ci, level, text))

        if not heading_map:
            return None, None, []

        # --- Pass 2: find match ---
        match_i = None
        match_level = None
        match_title = None
        for hi, (ci, level, text) in enumerate(heading_map):
            if query in text.lower():
                match_i = hi
                match_level = level
                match_title = text
                break

        if match_i is None:
            return None, None, [h[2] for h in heading_map]

        # --- Pass 3: find end boundary ---
        start_ci = heading_map[match_i][0]
        end_ci = len(body_children)
        for hi in range(match_i + 1, len(heading_map)):
            next_ci, next_level, _ = heading_map[hi]
            if next_level <= match_level:
                end_ci = next_ci
                break

        # --- Pass 4: collect text from start_ci to end_ci ---
        lines = []
        for child in body_children[start_ci:end_ci]:
            if child.tag.endswith('}p'):
                para = _DocxPara(child, doc)
                if para.text.strip():
                    lines.append(para.text)
            elif child.tag.endswith('}tbl'):
                # Include table cell text
                table = _DocxTable(child, doc)
                for row in table.rows:
                    row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_cells:
                        lines.append(' | '.join(row_cells))

        return '\n'.join(lines), match_title, [h[2] for h in heading_map]

    def _list_files_by_type(self, file_type: str = 'all', docs_dir: str = None) -> List[Dict]:
        """List files by type, scanning recursively into subfolders."""
        directory = docs_dir or self.docs_directory
        files = []
        try:
            for rel_path in storage.list_prefix(directory):
                extension = os.path.splitext(rel_path)[1].lower()
                if not extension:
                    continue  # skip bare directory entries
                if file_type == 'word' and extension not in ['.docx', '.doc']:
                    continue
                elif file_type == 'excel' and extension not in ['.xlsx', '.xls', '.xlsm']:
                    continue
                elif file_type == 'all' and extension not in ['.docx', '.doc', '.xlsx', '.xls', '.xlsm']:
                    continue
                abs_path = os.path.join(directory, rel_path)
                try:
                    import pathlib as _pl
                    stat = _pl.Path(abs_path).stat()
                    size = stat.st_size
                    modified = stat.st_mtime
                except Exception:
                    size = 0
                    modified = 0.0
                files.append({
                    'filename': rel_path,
                    'path': abs_path,
                    'extension': extension,
                    'size': size,
                    'modified': modified,
                })
            return sorted(files, key=lambda x: x['modified'], reverse=True)
        except Exception as e:
            self.logger.error(f"Failed to list files: {e}")
            return []

    def _read_word_document(self, file_path: str) -> Dict:
        """Read Word document via storage abstraction (REQ-PREP-03)."""
        try:
            from docx import Document

            raw = storage.read_bytes(file_path)
            doc = Document(io.BytesIO(raw))
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract text from tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                'filename': os.path.basename(file_path),
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs),
                'tables': tables,
                'table_count': len(tables)
            }
            
        except ImportError:
            raise ValueError("python-docx library not installed. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to read Word document: {str(e)}")

    def _read_excel_document(self, file_path: str, sheet_name: Optional[str] = None,
                            sheet_index: Optional[int] = None, max_rows: int = 100,
                            include_formulas: bool = False) -> Dict:
        """Read Excel document via storage abstraction (REQ-PREP-03)."""
        try:
            from openpyxl import load_workbook

            raw = storage.read_bytes(file_path)
            wb = load_workbook(io.BytesIO(raw), data_only=not include_formulas)
            
            # Get sheet
            if sheet_name:
                sheet = wb[sheet_name]
            elif sheet_index is not None:
                sheet = wb.worksheets[sheet_index]
            else:
                sheet = wb.active
            
            # Extract data
            data = []
            for i, row in enumerate(sheet.iter_rows(values_only=not include_formulas)):
                if i >= max_rows:
                    break
                data.append(list(row))
            
            result = {
                'filename': os.path.basename(file_path),
                'sheet_name': sheet.title,
                'data': data,
                'row_count': len(data),
                'column_count': len(data[0]) if data else 0
            }
            
            if include_formulas:
                formulas = []
                for row in sheet.iter_rows():
                    row_formulas = []
                    for cell in row:
                        if cell.value and isinstance(cell.value, str) and cell.value.startswith('='):
                            row_formulas.append({
                                'cell': cell.coordinate,
                                'formula': cell.value
                            })
                    if row_formulas:
                        formulas.append(row_formulas)
                result['formulas'] = formulas
            
            wb.close()
            return result
            
        except ImportError:
            raise ValueError("openpyxl library not installed. Install with: pip install openpyxl")
        except Exception as e:
            raise ValueError(f"Failed to read Excel document: {str(e)}")


class MsDocListFilesTool(MsDocBaseTool):
    """
    Tool to list available files in the docs directory
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_list_files',
            'description': 'List all Word and Excel documents in the docs directory',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "file_type": {
                    "type": "string",
                    "description": "Type of files to list",
                    "enum": ["all", "word", "excel"],
                    "default": "all"
                },
                "section": _SECTION_PROP,
            }
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "directory": {"type": "string"},
                "file_type": {"type": "string"},
                "count": {"type": "integer"},
                "files": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "filename": {"type": "string"},
                            "path": {"type": "string"},
                            "extension": {"type": "string"},
                            "size": {"type": "integer"},
                            "modified": {"type": "number"}
                        }
                    }
                }
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute list files — scans all data layers when section is 'domain_data'
        so documents are found regardless of which subfolder they were uploaded to."""
        section = arguments.get('section', 'domain_data')
        file_type = arguments.get('file_type', 'all')

        # When listing domain_data, scan the full domain_data tree (recursive)
        # plus my_data and common — so files in any subfolder (e.g. "bmo financials/")
        # are discovered without needing to know the subfolder name in advance.
        if section == 'domain_data':
            dirs_to_scan = []
            try:
                from flask import g as _g
                r = getattr(_g, 'worker_data_root', None)
                if r:
                    dirs_to_scan.append(r.rstrip('/'))
                r2 = getattr(_g, 'worker_my_data_root', None)
                if r2:
                    dirs_to_scan.append(r2.rstrip('/'))
                r3 = getattr(_g, 'worker_common_root', None)
                if r3:
                    dirs_to_scan.append(r3.rstrip('/'))
            except RuntimeError:
                pass
            if not dirs_to_scan:
                dirs_to_scan = [self.docs_directory]

            seen = set()
            all_files = []
            for d in dirs_to_scan:
                for f in self._list_files_by_type(file_type, d):
                    # Use subfolder-relative filename as dedup key so
                    # "bmo financials/Suppq425.xlsx" and "Suppq425.xlsx" are distinct
                    key = f['path']
                    if key not in seen:
                        seen.add(key)
                        all_files.append(f)

            return {
                'directory': dirs_to_scan[0],
                'section': section,
                'file_type': file_type,
                'count': len(all_files),
                'files': all_files,
                '_source': ', '.join(dirs_to_scan),
            }

        # For my_data / common — single directory as before
        docs_dir = self._resolve_docs_dir(section)
        files = self._list_files_by_type(file_type, docs_dir)
        return {
            'directory': docs_dir,
            'section': section,
            'file_type': file_type,
            'count': len(files),
            'files': files,
            '_source': docs_dir
        }


class MsDocReadWordTool(MsDocBaseTool):
    """
    Tool to read Word documents
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_read_word',
            'description': 'Read and extract content from Word documents (.docx)',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Word file to read"
                },
                "heading": {
                    "type": "string",
                    "description": (
                        "For large Word documents: extract only the section under this heading. "
                        "Case-insensitive partial match against Word Heading styles (Heading 1/2/3). "
                        "E.g. 'Market Risk', 'Executive Summary', 'Introduction'. "
                        "If not found, returns available_headings[] so you can pick the right one."
                    ),
                },
                "section": _SECTION_PROP,
            },
            "required": ["filename"]
        }

    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "paragraphs": {
                    "type": "array",
                    "items": {"type": "string"}
                },
                "paragraph_count": {"type": "integer"},
                "tables": {
                    "type": "array",
                    "items": {
                        "type": "array",
                        "items": {
                            "type": "array",
                            "items": {"type": "string"}
                        }
                    }
                },
                "table_count": {"type": "integer"}
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute read Word document"""
        filename = arguments['filename']
        section = arguments.get('section', 'domain_data')
        heading = (arguments.get('heading') or '').strip()
        file_path = self._get_file_path(filename, section)

        if not storage.exists(file_path):
            raise ValueError(f"File not found: {filename}")

        if heading:
            from docx import Document
            raw = storage.read_bytes(file_path)
            doc = Document(io.BytesIO(raw))
            text, matched_title, all_headings = self._extract_word_heading(doc, heading)

            if text is None:
                if not all_headings:
                    return {
                        'error': (
                            f'No Word heading styles found in {filename}. '
                            'The document may not use standard Heading styles. '
                            'Try msdoc_read_word without heading= to read the full document.'
                        ),
                        'filename': filename,
                        '_source': file_path,
                    }
                return {
                    'error': f'Heading "{heading}" not found in {filename}',
                    'available_headings': all_headings[:40],
                    'filename': filename,
                    '_source': file_path,
                }

            return {
                'filename': filename,
                'matched_heading': matched_title,
                'size_chars': len(text),
                'content': text,
                'available_headings': all_headings[:40],
                '_source': file_path,
            }

        result = self._read_word_document(file_path)
        result['_source'] = str(file_path)
        return result


class MsDocReadExcelTool(MsDocBaseTool):
    """
    Tool to read Excel spreadsheets
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_read_excel',
            'description': 'Read and extract data from Excel spreadsheets (.xlsx)',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Excel file to read"
                },
                "sheet_name": {
                    "type": "string",
                    "description": "Name of the sheet to read (optional)"
                },
                "sheet_index": {
                    "type": "integer",
                    "description": "Index of the sheet to read (0-based, optional)",
                    "minimum": 0
                },
                "max_rows": {
                    "type": "integer",
                    "description": "Maximum number of rows to return",
                    "default": 100,
                    "minimum": 1,
                    "maximum": 10000
                },
                "include_formulas": {
                    "type": "boolean",
                    "description": "Include cell formulas",
                    "default": False
                },
                "section": _SECTION_PROP,
            },
            "required": ["filename"]
        }

    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "sheet_name": {"type": "string"},
                "data": {
                    "type": "array",
                    "items": {
                        "type": "array"
                    }
                },
                "row_count": {"type": "integer"},
                "column_count": {"type": "integer"},
                "formulas": {
                    "type": "array",
                    "items": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "cell": {"type": "string"},
                                "formula": {"type": "string"}
                            }
                        }
                    }
                }
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute read Excel document"""
        filename = arguments['filename']
        section = arguments.get('section', 'domain_data')
        file_path = self._get_file_path(filename, section)

        if not storage.exists(file_path):
            raise ValueError(f"File not found: {filename}")

        result = self._read_excel_document(
            file_path,
            arguments.get('sheet_name'),
            arguments.get('sheet_index'),
            arguments.get('max_rows', 100),
            arguments.get('include_formulas', False)
        )
        result['_source'] = str(file_path)
        return result


class MsDocSearchWordTool(MsDocBaseTool):
    """
    Tool to search for text in Word documents
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_search_word',
            'description': 'Search for text within Word documents',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Word file to search"
                },
                "search_term": {
                    "type": "string",
                    "description": "Text to search for"
                },
                "section": _SECTION_PROP,
            },
            "required": ["filename", "search_term"]
        }

    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "search_term": {"type": "string"},
                "matches": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "paragraph_index": {"type": "integer"},
                            "text": {"type": "string"}
                        }
                    }
                },
                "match_count": {"type": "integer"}
            }
        }

    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute search in Word document"""
        filename = arguments['filename']
        search_term = arguments['search_term'].lower()
        section = arguments.get('section', 'domain_data')
        file_path = self._get_file_path(filename, section)
        
        if not storage.exists(file_path):
            raise ValueError(f"File not found: {filename}")
        
        doc_content = self._read_word_document(file_path)
        
        matches = []
        for i, paragraph in enumerate(doc_content['paragraphs']):
            if search_term in paragraph.lower():
                matches.append({
                    'paragraph_index': i,
                    'text': paragraph
                })
        
        return {
            'filename': filename,
            'search_term': search_term,
            'matches': matches,
            'match_count': len(matches),
            '_source': str(file_path)
        }


class MsDocSearchExcelTool(MsDocBaseTool):
    """
    Tool to search for text in Excel spreadsheets
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_search_excel',
            'description': 'Search for text within Excel spreadsheets',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Excel file to search"
                },
                "search_term": {
                    "type": "string",
                    "description": "Text to search for"
                },
                "sheet_name": {
                    "type": "string",
                    "description": "Name of the sheet to search (optional)"
                },
                "section": _SECTION_PROP,
            },
            "required": ["filename", "search_term"]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "sheet_name": {"type": "string"},
                "search_term": {"type": "string"},
                "matches": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "row_index": {"type": "integer"},
                            "column_index": {"type": "integer"},
                            "value": {"type": "string"}
                        }
                    }
                },
                "match_count": {"type": "integer"}
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute search in Excel document"""
        filename = arguments['filename']
        search_term = arguments['search_term'].lower()
        sheet_name = arguments.get('sheet_name')
        section = arguments.get('section', 'domain_data')
        file_path = self._get_file_path(filename, section)
        
        if not storage.exists(file_path):
            raise ValueError(f"File not found: {filename}")
        
        excel_content = self._read_excel_document(
            file_path,
            sheet_name=sheet_name,
            max_rows=10000
        )
        
        matches = []
        for row_idx, row in enumerate(excel_content['data']):
            for col_idx, cell in enumerate(row):
                if cell and search_term in str(cell).lower():
                    matches.append({
                        'row_index': row_idx,
                        'column_index': col_idx,
                        'value': str(cell)
                    })
        
        return {
            'filename': filename,
            'sheet_name': excel_content['sheet_name'],
            'search_term': search_term,
            'matches': matches,
            'match_count': len(matches),
            '_source': str(file_path)
        }


class MsDocGetWordMetadataTool(MsDocBaseTool):
    """
    Tool to get Word document metadata
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_get_word_metadata',
            'description': 'Get metadata and properties from Word documents',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Word file"
                },
                "section": _SECTION_PROP,
            },
            "required": ["filename"]
        }

    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "metadata": {
                    "type": "object",
                    "properties": {
                        "author": {"type": "string"},
                        "title": {"type": "string"},
                        "subject": {"type": "string"},
                        "created": {"type": "string"},
                        "modified": {"type": "string"}
                    }
                }
            }
        }

    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute get Word metadata"""
        filename = arguments['filename']
        section = arguments.get('section', 'domain_data')
        file_path = self._get_file_path(filename, section)
        
        if not storage.exists(file_path):
            raise ValueError(f"File not found: {filename}")
        
        try:
            from docx import Document
            
            raw = storage.read_bytes(file_path)
            doc = Document(io.BytesIO(raw))
            core_props = doc.core_properties
            
            return {
                'filename': filename,
                'metadata': {
                    'author': core_props.author or '',
                    'title': core_props.title or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else ''
                },
                '_source': str(file_path)
            }
            
        except ImportError:
            raise ValueError("python-docx library not installed")
        except Exception as e:
            raise ValueError(f"Failed to get metadata: {str(e)}")


class MsDocGetExcelMetadataTool(MsDocBaseTool):
    """
    Tool to get Excel spreadsheet metadata
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_get_excel_metadata',
            'description': 'Get metadata and properties from Excel spreadsheets',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Excel file"
                },
                "section": _SECTION_PROP,
            },
            "required": ["filename"]
        }

    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "metadata": {
                    "type": "object",
                    "properties": {
                        "creator": {"type": "string"},
                        "title": {"type": "string"},
                        "subject": {"type": "string"},
                        "created": {"type": "string"},
                        "modified": {"type": "string"}
                    }
                }
            }
        }

    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute get Excel metadata"""
        filename = arguments['filename']
        section = arguments.get('section', 'domain_data')
        file_path = self._get_file_path(filename, section)

        if not storage.exists(file_path):
            raise ValueError(f"File not found: {filename}")

        try:
            from openpyxl import load_workbook

            raw = storage.read_bytes(file_path)
            wb = load_workbook(io.BytesIO(raw))
            props = wb.properties
            
            metadata = {
                'creator': props.creator or '',
                'title': props.title or '',
                'subject': props.subject or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else ''
            }
            
            wb.close()
            
            return {
                'filename': filename,
                'metadata': metadata,
                '_source': str(file_path)
            }

        except ImportError:
            raise ValueError("openpyxl library not installed")
        except Exception as e:
            raise ValueError(f"Failed to get metadata: {str(e)}")


class MsDocExtractTextTool(MsDocBaseTool):
    """
    Tool to extract plain text from documents
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_extract_text',
            'description': 'Extract all text content from Word or Excel documents',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the file to extract text from"
                },
                "section": _SECTION_PROP,
            },
            "required": ["filename"]
        }

    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "text": {"type": "string"},
                "character_count": {"type": "integer"}
            }
        }

    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute extract text"""
        filename = arguments['filename']
        section = arguments.get('section', 'domain_data')
        file_path = self._get_file_path(filename, section)

        if not storage.exists(file_path):
            raise ValueError(f"File not found: {filename}")

        extension = os.path.splitext(file_path)[1].lower()
        
        if extension in ['.docx', '.doc']:
            doc_content = self._read_word_document(file_path)
            text = '\n'.join(doc_content['paragraphs'])
        elif extension in ['.xlsx', '.xls', '.xlsm']:
            excel_content = self._read_excel_document(file_path, max_rows=10000)
            text_parts = []
            for row in excel_content['data']:
                text_parts.append('\t'.join(str(cell) if cell else '' for cell in row))
            text = '\n'.join(text_parts)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
        
        return {
            'filename': filename,
            'text': text,
            'character_count': len(text),
            '_source': str(file_path)
        }


class MsDocGetExcelSheetsTool(MsDocBaseTool):
    """
    Tool to list all sheets in an Excel workbook
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_get_excel_sheets',
            'description': 'List all sheets in an Excel workbook',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Excel file"
                },
                "section": _SECTION_PROP,
            },
            "required": ["filename"]
        }

    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "sheets": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "index": {"type": "integer"},
                            "name": {"type": "string"}
                        }
                    }
                },
                "count": {"type": "integer"}
            }
        }

    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute get Excel sheets"""
        filename = arguments['filename']
        section = arguments.get('section', 'domain_data')
        file_path = self._get_file_path(filename, section)

        if not storage.exists(file_path):
            raise ValueError(f"File not found: {filename}")

        try:
            from openpyxl import load_workbook

            raw = storage.read_bytes(file_path)
            wb = load_workbook(io.BytesIO(raw), read_only=True)
            sheets = [
                {'index': i, 'name': sheet.title}
                for i, sheet in enumerate(wb.worksheets)
            ]
            wb.close()
            
            return {
                'filename': filename,
                'sheets': sheets,
                'count': len(sheets),
                '_source': str(file_path)
            }
            
        except ImportError:
            raise ValueError("openpyxl library not installed")
        except Exception as e:
            raise ValueError(f"Failed to get sheets: {str(e)}")


class MsDocReadExcelSheetTool(MsDocBaseTool):
    """
    Tool to read a specific sheet from an Excel workbook
    """
    
    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_read_excel_sheet',
            'description': 'Read data from a specific sheet in an Excel workbook',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
    
    def get_input_schema(self) -> Dict:
        """Get input schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name of the Excel file"
                },
                "sheet_name": {
                    "type": "string",
                    "description": "Name of the sheet to read"
                },
                "sheet_index": {
                    "type": "integer",
                    "description": "Index of the sheet to read (0-based)",
                    "minimum": 0
                },
                "max_rows": {
                    "type": "integer",
                    "description": "Maximum number of rows to return",
                    "default": 100,
                    "minimum": 1,
                    "maximum": 10000
                },
                "section": _SECTION_PROP,
            },
            "required": ["filename"]
        }
    
    def get_output_schema(self) -> Dict:
        """Get output schema"""
        return {
            "type": "object",
            "properties": {
                "filename": {"type": "string"},
                "sheet_name": {"type": "string"},
                "data": {
                    "type": "array",
                    "items": {
                        "type": "array"
                    }
                },
                "row_count": {"type": "integer"},
                "column_count": {"type": "integer"}
            }
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Dict:
        """Execute read Excel sheet"""
        filename = arguments['filename']
        section = arguments.get('section', 'domain_data')
        file_path = self._get_file_path(filename, section)

        if not storage.exists(file_path):
            raise ValueError(f"File not found: {filename}")

        result = self._read_excel_document(
            file_path,
            arguments.get('sheet_name'),
            arguments.get('sheet_index'),
            arguments.get('max_rows', 100),
            False
        )
        result['_source'] = str(file_path)
        return result


class MsDocGetExcelStatsTool(MsDocBaseTool):
    """
    Tool to get per-sheet statistics for an Excel workbook.
    Returns row/column counts, detected data types, and numeric summary stats.
    """

    def __init__(self, config: Dict = None):
        default_config = {
            'name': 'msdoc_get_excel_stats',
            'description': (
                'Get per-sheet statistics for an Excel workbook: row/column counts, '
                'column names, data types, and numeric summary (min/max/mean) for each sheet.'
            ),
            'version': '1.0.0',
            'enabled': True,
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)

    def get_input_schema(self) -> Dict:
        return {
            'type': 'object',
            'properties': {
                'filename': {
                    'type': 'string',
                    'description': 'Name of the Excel file',
                },
                'section': _SECTION_PROP,
            },
            'required': ['filename'],
        }

    def get_output_schema(self) -> Dict:
        return {'type': 'object'}

    def execute(self, arguments: Dict[str, Any]) -> Dict:
        filename = arguments['filename']
        section = arguments.get('section', 'domain_data')
        file_path = self._get_file_path(filename, section)

        if not storage.exists(file_path):
            raise ValueError(f'File not found: {filename}')

        try:
            from openpyxl import load_workbook

            raw = storage.read_bytes(file_path)
            wb = load_workbook(io.BytesIO(raw), data_only=True)
            sheets_stats = []

            for sheet in wb.worksheets:
                rows = list(sheet.iter_rows(values_only=True))
                if not rows:
                    sheets_stats.append({
                        'sheet': sheet.title,
                        'row_count': 0,
                        'column_count': 0,
                        'columns': [],
                        'numeric_summary': {},
                    })
                    continue

                header = [str(c) if c is not None else f'col_{i}' for i, c in enumerate(rows[0])]
                data_rows = rows[1:]
                col_count = len(header)

                # Per-column type detection and numeric summary
                col_stats = {}
                for ci, col_name in enumerate(header):
                    values = [r[ci] for r in data_rows if ci < len(r) and r[ci] is not None]
                    nums = []
                    for v in values:
                        try:
                            nums.append(float(v))
                        except (TypeError, ValueError):
                            pass
                    if nums:
                        col_stats[col_name] = {
                            'type': 'numeric',
                            'non_null': len(values),
                            'min': round(min(nums), 4),
                            'max': round(max(nums), 4),
                            'mean': round(sum(nums) / len(nums), 4),
                        }
                    else:
                        col_stats[col_name] = {
                            'type': 'text',
                            'non_null': len(values),
                        }

                sheets_stats.append({
                    'sheet': sheet.title,
                    'row_count': len(data_rows),
                    'column_count': col_count,
                    'columns': header,
                    'column_stats': col_stats,
                })

            wb.close()
            return {
                'filename': filename,
                'sheet_count': len(sheets_stats),
                'sheets': sheets_stats,
                '_source': file_path,
            }

        except ImportError:
            raise ValueError('openpyxl library not installed. Install with: pip install openpyxl')
        except Exception as e:
            raise ValueError(f'Failed to get Excel stats: {str(e)}')


# Tool registry
MSDOC_TOOLS = {
    'msdoc_list_files': MsDocListFilesTool,
    'msdoc_read_word': MsDocReadWordTool,
    'msdoc_read_excel': MsDocReadExcelTool,
    'msdoc_search_word': MsDocSearchWordTool,
    'msdoc_search_excel': MsDocSearchExcelTool,
    'msdoc_get_word_metadata': MsDocGetWordMetadataTool,
    'msdoc_get_excel_metadata': MsDocGetExcelMetadataTool,
    'msdoc_extract_text': MsDocExtractTextTool,
    'msdoc_get_excel_sheets': MsDocGetExcelSheetsTool,
    'msdoc_read_excel_sheet': MsDocReadExcelSheetTool,
    'msdoc_get_excel_stats': MsDocGetExcelStatsTool,
}
